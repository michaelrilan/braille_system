from django.contrib.auth import authenticate, login, logout, update_session_auth_hash
from django.shortcuts import render, redirect, reverse
from django.http import HttpResponse, request
from django.contrib.auth.password_validation import validate_password
from django.contrib.auth.decorators import login_required
from .models import *
from django.contrib import messages
from django.contrib.auth.models import User
from django.core.exceptions import ValidationError
import os
import json
from datetime import datetime
from happytransformer import HappyTextToText
from happytransformer import TTSettings
from textblob import TextBlob
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.cache import cache_control
from django.shortcuts import get_object_or_404
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import random
from django.utils.timezone import now
import string
from pydub import AudioSegment
import speech_recognition as sr
import socket
import logging
import shutil
from django.core.serializers.json import DjangoJSONEncoder
from django.utils.safestring import mark_safe
from django.db.models import Prefetch
logger = logging.getLogger(__name__)

def check_internet_connection():
    try:
        # Test a connection to a reliable service
        response = requests.get('https://www.google.com/', timeout=5)
        return response.status_code == 200
    except requests.ConnectionError as e:
        logger.error(f"Internet connection error: {e}")
        return False

def generate_random_string(length=8):
    # Define the possible characters
    characters = string.ascii_letters + string.digits
    
    # Generate a random string
    random_string = ''.join(random.choice(characters) for _ in range(length))
    
    return "#" + random_string



# Helper function to convert audio file to WAV format
def convert_to_wav(input_file_path, output_file_path):
    try:
        # Load the audio file
        audio = AudioSegment.from_file(input_file_path)
        # Export as WAV
        audio.export(output_file_path, format='wav')
        return True
    except Exception as e:
        logger.error(f'Conversion error: {e}')
        return False
    

@csrf_exempt
def upload_audio(request):
    if request.method == 'POST':
        audio_file = request.FILES.get('audio')
        
        if audio_file:
            logger.info(f'File uploaded: {audio_file.name}, Size: {audio_file.size}, Content-Type: {audio_file.content_type}')
            
            temp_input_path = 'temp_audio_input'
            temp_output_path = 'temp_audio_output.wav'
            
            try:
                with open(temp_input_path, 'wb') as temp_file:
                    for chunk in audio_file.chunks():
                        temp_file.write(chunk)
                
                if not convert_to_wav(temp_input_path, temp_output_path):
                    return JsonResponse({'transcription': 'Failed to convert audio file'}, status=500)
                
                recognizer = sr.Recognizer()
                
                try:
                    with sr.AudioFile(temp_output_path) as source:
                        audio = recognizer.record(source)
                    
                    transcription = recognizer.recognize_google(audio)
                    return JsonResponse({'transcription': transcription})
                
                except ValueError as e:
                    logger.error(f'Audio file reading error: {e}')
                    return JsonResponse({'transcription': 'Invalid audio file format'}, status=400)
            
            except Exception as e:
                logger.error(f'Unexpected error while handling the file: {e}')
                return JsonResponse({'transcription': 'Sorry I cant catch your voice at the moment, please try again later'}, status=500)
            
            finally:
                if os.path.exists(temp_input_path):
                    os.remove(temp_input_path)
                if os.path.exists(temp_output_path):
                    os.remove(temp_output_path)
        
        return JsonResponse({'transcription': 'No audio file received'}, status=400)
    
    return JsonResponse({'transcription': 'Invalid request'}, status=400)





def generate_unique_username(first_name, last_name, existing_usernames):
    # Split the first name to handle multiple words (e.g., "John Dave")
    first_name_parts = first_name.split()
    
    # Get the first letter of each part of the first name
    first_letters = ''.join([name[0].lower() for name in first_name_parts])
    
    # Lowercase the last name
    last_name_lower = last_name.lower()
    
    # Combine the first letters and the lowercase last name to form the initial username
    base_username = f"{first_letters}.{last_name_lower}"
    unique_username = base_username
    count = 1

    # Check if the username is unique, if not, append a number to make it unique
    while unique_username in existing_usernames:
        unique_username = f"{base_username}{count}"
        count += 1
    
    return unique_username



brailleDict = {
  'a': '⠁', 'b': '⠃', 'c': '⠉', 'd': '⠙', 'e': '⠑', 'f': '⠋', 'g': '⠛',
  'h': '⠓', 'i': '⠊', 'j': '⠚', 'k': '⠅', 'l': '⠇', 'm': '⠍', 'n': '⠝',
  'o': '⠕', 'p': '⠏', 'q': '⠟', 'r': '⠗', 's': '⠎', 't': '⠞', 'u': '⠥',
  'v': '⠧', 'w': '⠺', 'x': '⠭', 'y': '⠽', 'z': '⠵','A': '⠠⠁', 'B': '⠠⠃', 
  'C': '⠠⠉', 'D': '⠠⠙', 'E': '⠠⠑', 'F': '⠠⠋', 'G': '⠠⠛','H': '⠠⠓', 'I': '⠠⠊',
  'J': '⠠⠚', 'K': '⠠⠅', 'L': '⠠⠇', 'M': '⠠⠍', 'N': '⠠⠝','O': '⠠⠕', 'P': '⠠⠏',
  'Q': '⠠⠟', 'R': '⠠⠗', 'S': '⠠⠎', 'T': '⠠⠞', 'U': '⠠⠥',
  'V': '⠠⠧', 'W': '⠠⠺', 'X': '⠠⠭', 'Y': '⠠⠽', 'Z': '⠠⠵',
  '0': '⠴', '1': '⠂', '2': '⠆', '3': '⠒', '4': '⠲', '5': '⠢', '6': '⠖', '7': '⠶', '8': '⠦', '9': '⠔',
  ' ': ' ', 
  '.': '⠲',
  ',': '⠂', 
  '!': '⠖', 
  '?': '⠦', 
  "'": '⠄' 
}



def convert_to_braille(text):
    return ''.join(brailleDict.get(char, char) for char in text)



def uppercase(data):
    return str(data).upper()


def check_text(text):
    # Create a TextBlob object
    blob = TextBlob(text)
    
    # Correct spelling
    corrected_text = blob.correct()
    
    # Print the original and corrected text
    print("Original Text:", text)
    print("Corrected Text:", corrected_text)
    
    return(corrected_text)


        
@csrf_exempt
def grammar_check(request):
    if request.method == 'POST':
        data = request.POST.get('text', '')
        if not data:
            return JsonResponse({'error': 'No text provided'}, status=400)
        
        text = check_text(data)
        happy_tt = HappyTextToText("T5",  "prithivida/grammar_error_correcter_v1")
        settings = TTSettings(do_sample=True, top_k=10, temperature=0.5,  min_length=1, max_length=100)

        text = "gec: " + str(text)
        

        result = happy_tt.generate_text(text, args=settings)
        print('grammar_check: ' + str(result.text))
        
        return JsonResponse({'corrected_text': result.text})
    return JsonResponse({'error': 'Invalid request method'}, status=405)




@cache_control(no_cache=True, must_revalidate=True, no_store=True)
def user_login(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        
        if user is not None:

            current_year = now().year
            threshold_year = current_year - 5
            UserProfile.objects.filter(
                school_year__lt=str(threshold_year)
            ).update(deleteflag=True)

            # Check if the user has a profile and its deleteflag status
            user_profile = UserProfile.objects.get(user=user)
            if user_profile.deleteflag:
                messages.error(request, "Invalid Credentials!")
                return redirect('login')
            else:
                login(request, user)
                return redirect('dashboard')
        else:
            messages.error(request, "Invalid Credentials!")
    return render(request, 'login.html')





@cache_control(no_cache=True, must_revalidate=True, no_store=True)
def register(request):
    if request.method == 'POST':
        first_name =uppercase(request.POST.get('first_name')) 
        last_name = uppercase(request.POST.get('last_name'))
        email = request.POST.get('email')
        username = request.POST.get('username')
        pw1 = request.POST.get('password1')
        pw2 = request.POST.get('password2')

        email_check =  User.objects.filter(email = email).count()
        fname_check =  User.objects.filter(first_name = first_name,last_name = last_name).count()
        username_check = User.objects.filter(username = username).count()
        if email_check > 0:
            messages.error(request,'Email Already Exists')
        elif fname_check>0:
            messages.error(request,'Your Fullname Already Exists')
        elif username_check > 0:
            messages.error(request,'Username Already Exists, Try Another')
        elif pw1 != pw2:
            messages.error(request,'Passwords Not Equal')
        else:

            try:
                validate_password(pw1)
                user = User.objects.create_user(username=username, email=email, password=pw1, first_name=first_name, last_name=last_name)
                user.save()
                messages.success(request,'Registered Successfully')
                return redirect('login')
            except ValidationError as e:
                # Handle the validation error (e.g., log it, display a message)
                messages.error(request, f'Invalid: {", ".join(e.messages)}')

        context = {
            'first_name': first_name,
            'last_name': last_name,
            'email': email,
            'username': username
        }
        return render(request, "register.html", context)
    return render(request, 'register.html')




@cache_control(no_cache=True, must_revalidate=True, no_store=True)
@login_required(login_url='login')
def dashboard(request):
    if request.user.is_authenticated:
        user_id = request.user.id
        braille_infos_count = BrailleInfo.objects.filter(user_id=user_id).count()
        shared_to_me = SharedBraille.objects.filter(shared_to_user = user_id).count()
        archives = BrailleInfo.objects.filter(user_id = user_id,deleteflag=True).count()
        shared_braille_entries = SharedBraille.objects.select_related('user', 'shared_to_user', 'braille_info').filter(user_id=user_id).count()

        all_shared_entries = SharedBraille.objects.filter(user_id=user_id).select_related('braille_info', 'shared_to_user')
        # Filter to retain only distinct braille_info_id entries
        seen_braille_ids = set()
        distinct_shared_entries = []
        for entry in all_shared_entries:
            if entry.braille_info_id not in seen_braille_ids:
                seen_braille_ids.add(entry.braille_info_id)
                distinct_shared_entries.append(entry)

        student_dashboard_tbl = SharedBraille.objects.filter(shared_to_user = user_id, is_viewed = False).all()
                
        context = {
            'student_dashboard_tbl': student_dashboard_tbl,
            'admin_dashboard_tbl': distinct_shared_entries,
            'archives_count' : archives,
            'braille_infos_count': braille_infos_count,
            'shared_to_me_count': shared_to_me,
            'shared_braille_entries': shared_braille_entries,
        }
    return render(request, 'dashboard.html',context)


@cache_control(no_cache=True, must_revalidate=True, no_store=True)
@login_required(login_url='login')
def notif_braille(request,shared_id):
    shared_braille = get_object_or_404(SharedBraille, id=shared_id)
    shared_braille.is_viewed = True
    shared_braille.save()
    if request.method == 'POST':
            braille_id = request.POST.get('braille_id')
            filename = request.POST.get('filename')
            documents_dir = 'static/documents'
            
            fetch_braille = BrailleInfo.objects.get(id=braille_id, deleteflag=False)
            title = fetch_braille.title
            braille_text = fetch_braille.braille_text
            braille_draft = fetch_braille.braille_draft

            # Create the documents directory if it doesn't exist
            if not os.path.exists(documents_dir):
                os.makedirs(documents_dir)

            file_path = os.path.join(documents_dir, filename)

            # Check if the file exists, if not, recreate it
            if not os.path.exists(file_path):
                try:
                    # Create and save the document
                    # document = Document()
                    # document.add_heading(title, level=1)
                    # document.add_paragraph(braille_text)
                    # document.add_paragraph(braille_draft)
                    # document.save(file_path)

                    # Create a new document
                    document = Document()

                    # Add title with font size 15 and center alignment
                    title_paragraph = document.add_heading(title, level=1)
                    title_run = title_paragraph.runs[0]
                    title_run.font.size = Pt(15)
                    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centers the title

                    # Add space below the title (1 line break)
                    for _ in range(1):
                        document.add_paragraph()  # Add empty paragraphs for spacing

                    # Modify 'braille_text' to add 3 spaces between each character
                    braille_text_spaced = '   '.join(braille_text)  # Adds 3 spaces between each character

                    # Add braille_text with font size 22
                    para_braille = document.add_paragraph()
                    run_braille = para_braille.add_run(braille_text_spaced)
                    font_braille = run_braille.font
                    font_braille.size = Pt(22)

                    # Add braille_draft with font size 22
                    para_draft = document.add_paragraph()
                    run_draft = para_draft.add_run(braille_draft)
                    font_draft = run_draft.font
                    font_draft.size = Pt(22)

                    # Define the file path and save the document
                    file_path = os.path.join(documents_dir, filename)
                    document.save(file_path)
                except Exception as e:
                    print(f"Error creating document: {e}")

            try:
                # Determine the user's Downloads directory
                downloads_dir = os.path.join(os.path.expanduser("~"), "Downloads")
                destination_path = os.path.join(downloads_dir, filename)

                # Copy the file to the Downloads directory
                shutil.copy(file_path, destination_path)
                print(f"File copied to: {destination_path}")
                messages.success(request,"File Downloaded Successfully")
                # Optionally, you can also return a success message or redirect
                return redirect('notif_braile')  # Adjust this to your success view
            except Exception as e:
                print(f"Error during file copy: {e}")
    braille_infos =  SharedBraille.objects.filter(id=shared_id).all()
    
    context = {
        'braille_infos':braille_infos,
    }
    return render(request,'notif_braille.html',context)



@cache_control(no_cache=True, must_revalidate=True, no_store=True)
@login_required(login_url='login')
def notif_braille_admin(request, braille_id):
    user_id = request.user.id
    # shared = SharedBraille.objects.filter(braille_info = braille_id,user = user_id).all()
    shared = SharedBraille.objects.filter(
    braille_info=braille_id, user=user_id
    ).select_related('shared_to_user').prefetch_related(
        Prefetch(
            'shared_to_user__userprofile_set',
            queryset=UserProfile.objects.all()
        )
    )

    context = {
        "shared": shared
    }
    return render(request,'notif_braille_admin.html',context)


@cache_control(no_cache=True, must_revalidate=True, no_store=True)
@login_required(login_url='login')
def tutorial(request):
    if request.user.is_authenticated:
        user_id = request.user.id
        braille_infos_count = BrailleInfo.objects.filter(user_id=user_id).count()
        activity_history = ActivityHistory.objects.filter(user_id=user_id).order_by('-date_log')

        context = {
            'braille_infos_count': braille_infos_count,
            'activity_history':activity_history
        }
    return render(request, 'tutorial.html',context)



import os
from django.http import FileResponse
from django.conf import settings

def download_confirm(request):
    filename = 'fff.docx'  # Replace with your actual file name
    file_path = os.path.join(settings.BASE_DIR, 'static', 'documents', filename)

    # Check if file exists
    if os.path.exists(file_path):
        return FileResponse(open(file_path, 'rb'), as_attachment=True, filename=filename)
    else:
        # Handle case where file does not exist
        return render(request, 'download_confirm.html')


@cache_control(no_cache=True, must_revalidate=True, no_store=True)
@login_required(login_url='login')
def create_braille(request):
    current_date = datetime.now().strftime('%Y%m%d_%H%M%S')  # Format the date as YYYYMMDD
    if request.user.is_authenticated:
       if request.method == 'POST':
            title = request.POST.get('title')
            user_id = request.user.id
            braille_draft = request.POST.get('braille_draft')
            if (braille_draft.strip() == '' or title.strip() == ''):
                messages.error(request, 'Invalid content!')
                return redirect('create_braille')
            else:
                filename = f'{request.user.id}_{title.replace(" ", "_")}_{current_date}.docx'
                braille_text = convert_to_braille(braille_draft)
                braille_instance = BrailleInfo.objects.create(
                        user=request.user,
                        title=title,
                        braille_draft=braille_draft,
                        braille_text=braille_text,
                        filename = filename
                    )
                created_id = braille_instance.id
                # Ensure the directory exists
                # documents_dir = 'static/documents'
                # if not os.path.exists(documents_dir):
                #     os.makedirs(documents_dir)

                # document = Document()
                # document.add_heading(title, level=1)
                # document.add_paragraph(braille_text)
                # document.add_paragraph(braille_draft)

                # file_path = os.path.join('static/documents', filename)
                # document.save(file_path)


                # Set up the directory for saving documents
                documents_dir = 'static/documents'
                if not os.path.exists(documents_dir):
                    os.makedirs(documents_dir)

                # Create a new document
                document = Document()

                # Add title with font size 15 and center alignment
                title_paragraph = document.add_heading(title, level=1)
                title_run = title_paragraph.runs[0]
                title_run.font.size = Pt(15)
                title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centers the title

                # Add space below the title (1 line break)
                for _ in range(1):
                    document.add_paragraph()  # Add empty paragraphs for spacing

                # Modify 'braille_text' to add 3 spaces between each character
                braille_text_spaced = '   '.join(braille_text)  # Adds 3 spaces between each character

                # Add braille_text with font size 22
                para_braille = document.add_paragraph()
                run_braille = para_braille.add_run(braille_text_spaced)
                font_braille = run_braille.font
                font_braille.size = Pt(22)

                # Add braille_draft with font size 22
                para_draft = document.add_paragraph()
                run_draft = para_draft.add_run(braille_draft)
                font_draft = run_draft.font
                font_draft.size = Pt(22)

                # Define the file path and save the document
                file_path = os.path.join(documents_dir, filename)
                document.save(file_path)


                activity_history = ActivityHistory(user_id = user_id,activity_log="Created a New Braille File(File # " +str(created_id) + ")")
                activity_history.save()

                messages.success(request, 'Braille Successfully Created!')

                return redirect('create_braille')
    context = {'check_internet': check_internet_connection()}
    if not context['check_internet']:
        messages.info(request, 'No internet connection detected. Features like Grammar Checker and Speech to Text may not work')
    return render(request, 'create_braille.html', context)



@cache_control(no_cache=True, must_revalidate=True, no_store=True)
@login_required(login_url='login')
def view_braille(request):
    current_date = datetime.now().strftime('%Y%m%d_%H%M%S')  # Format the date as YYYYMMDD
    
    if request.user.is_authenticated:
        unique_school_years = UserProfile.objects.filter(
        is_active=True,
        is_student = True,
        deleteflag = False
        ).values_list('school_year', flat=True).distinct()
        
        print(unique_school_years)
        user_id = request.user.id
        user_profile = get_object_or_404(UserProfile, user_id=user_id)
        if request.method == 'POST':
            form_type = request.POST.get('form_type')

            # DOWNLOAD BRAILLE FUNCTION
            # if form_type == 'download_braille':
            #     print('download')
            #     braille_id = request.POST.get('braille_id')
            #     filename = request.POST.get('filename')
            #     documents_dir = 'static/documents'
            #     fetch_braille = BrailleInfo.objects.get(id=braille_id,deleteflag = False)
            #     title = fetch_braille.title
            #     braille_text = fetch_braille.braille_text
            #     braille_draft = fetch_braille.braille_draft

            #     if not os.path.exists(documents_dir):
            #         os.makedirs(documents_dir)
            #     file_path = os.path.join(documents_dir, filename)
            #     # Check if the file exists, if not, recreate it
            #     if not os.path.exists(file_path):
            #         try:
            #             # Create and save the document
            #             document = Document()
            #             document.add_heading(title, level=1)
            #             document.add_paragraph(braille_text)
            #             document.add_paragraph(braille_draft)
            #             document.save(file_path)
            #         except Exception as e:
            #             print(f"Error creating document: {e}")
                
            #     try:
                    
            #         return redirect('download_braille', file_name= filename) 
                
            #     except Exception as e:
            #         print(f"Error during download redirection: {e}")
            if form_type == 'download_braille':
                print('download')
                braille_id = request.POST.get('braille_id')
                filename = request.POST.get('filename')
                documents_dir = 'static/documents'
                
                fetch_braille = BrailleInfo.objects.get(id=braille_id, deleteflag=False)
                title = fetch_braille.title
                braille_text = fetch_braille.braille_text
                braille_draft = fetch_braille.braille_draft

                # Create the documents directory if it doesn't exist
                if not os.path.exists(documents_dir):
                    os.makedirs(documents_dir)

                file_path = os.path.join(documents_dir, filename)

                # Check if the file exists, if not, recreate it
                if not os.path.exists(file_path):
                    try:
                        # Create a new document
                        document = Document()

                        # Add title with font size 15 and center alignment
                        title_paragraph = document.add_heading(title, level=1)
                        title_run = title_paragraph.runs[0]
                        title_run.font.size = Pt(15)
                        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centers the title

                        # Add space below the title (1 line break)
                        for _ in range(1):
                            document.add_paragraph()  # Add empty paragraphs for spacing

                        # Modify 'braille_text' to add 3 spaces between each character
                        braille_text_spaced = '   '.join(braille_text)  # Adds 3 spaces between each character

                        # Add braille_text with font size 22
                        para_braille = document.add_paragraph()
                        run_braille = para_braille.add_run(braille_text_spaced)
                        font_braille = run_braille.font
                        font_braille.size = Pt(22)

                        # Add braille_draft with font size 22
                        para_draft = document.add_paragraph()
                        run_draft = para_draft.add_run(braille_draft)
                        font_draft = run_draft.font
                        font_draft.size = Pt(22)

                        # Define the file path and save the document
                        file_path = os.path.join(documents_dir, filename)
                        document.save(file_path)
                    except Exception as e:
                        print(f"Error creating document: {e}")

                try:
                    # Determine the user's Downloads directory
                    downloads_dir = os.path.join(os.path.expanduser("~"), "Downloads")
                    destination_path = os.path.join(downloads_dir, filename)

                    # Copy the file to the Downloads directory
                    shutil.copy(file_path, destination_path)
                    print(f"File copied to: {destination_path}")
                    messages.success(request,"File Downloaded Successfully")
                    # Optionally, you can also return a success message or redirect
                    return redirect('view_braille')  # Adjust this to your success view
                except Exception as e:
                    print(f"Error during file copy: {e}")
                    

            # EDIT BRAILLE FUNCTION
            
            if form_type == 'edit_braille':
                braille_id = request.POST.get('braille_id')
                braille_draft = request.POST.get('braille_draft')
                braille_text = convert_to_braille(braille_draft)
                title = request.POST.get('title')
                braille_info = BrailleInfo.objects.get(id=braille_id, deleteflag=False)

                document = Document()

                # Add title with font size 15 and center alignment
                title_paragraph = document.add_heading(title, level=1)
                title_run = title_paragraph.runs[0]
                title_run.font.size = Pt(15)
                title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centers the title

                # Add space below the title (3 line breaks)
                for _ in range(3):
                    document.add_paragraph()  # Add empty paragraphs for spacing

                # Modify 'braille_text' to add 3 spaces between each character
                braille_text_spaced = '   '.join(braille_text)  # Adds 3 spaces between each character

                # Add braille_text with font size 22
                para_braille = document.add_paragraph()
                run_braille = para_braille.add_run(braille_text_spaced)
                font_braille = run_braille.font
                font_braille.size = Pt(22)

                # Add braille_draft with font size 22
                para_draft = document.add_paragraph()
                run_draft = para_draft.add_run(braille_draft)
                font_draft = run_draft.font
                font_draft.size = Pt(22)

                # Define the filename and save the document
                filename = f'{user_id}_{title.replace(" ", "_")}_{current_date}.docx'
                file_path = os.path.join('static/documents', filename)
                document.save(file_path)


                # braille_id = request.POST.get('braille_id')
                # braille_draft = request.POST.get('braille_draft')
                # braille_text = convert_to_braille(braille_draft)
                # title = request.POST.get('title')
                # braille_info = BrailleInfo.objects.get(id=braille_id, deleteflag=False)
                
                # document = Document()
                # document.add_heading(title, level=1)
                # document.add_paragraph(braille_text)
                # document.add_paragraph(braille_draft)
                # filename = f'{user_id}_{title.replace(" ", "_")}_{current_date}.docx'
                # file_path = os.path.join('static/documents', filename)
                # document.save(file_path)

                braille_info.filename = filename
                braille_info.braille_draft = braille_draft
                braille_info.braille_text = braille_text
                braille_info.title = title
                braille_info.save()

                created_id = braille_info.id
                activity_history = ActivityHistory(user_id = user_id,activity_log="Updated Braille File(File # " +str(created_id) + ")")
                activity_history.save()
                messages.success(request, 'Braille Information Updated')
                return redirect('view_braille')

            elif form_type == 'archive_braille':
                braille_id = request.POST.get('braille_id')
                braille_info = BrailleInfo.objects.get(id=braille_id, deleteflag=False)
                braille_info.deleteflag = True
                braille_info.save()
                messages.success(request, 'Braille Successfully moved to Archive')
                return redirect('view_braille')
            

            elif form_type == 'share_braille':
                braille_id = request.POST.get('braille_id')
                school_year = request.POST.get('school_year')
                print(school_year)
                user_ids = UserProfile.objects.filter(
                deleteflag=False,
                is_active=True,
                is_student = True,
                school_year = school_year
                ).values_list('user_id', flat=True)

                user_ids_list = list(user_ids)
                any_new_entry_created = False
                braille_info = BrailleInfo.objects.get(id=braille_id)
                for shared_to_user_id in user_ids_list:
                    shared_to_user = User.objects.get(id=shared_to_user_id)
                    exists = SharedBraille.objects.filter(
                    braille_info=braille_info,
                    shared_to_user=shared_to_user,
                    user_id=user_id
                    ).exists()

                    if not exists:
                        SharedBraille.objects.create(
                            user_id=user_id,
                            shared_to_user=shared_to_user,
                            braille_info=braille_info
                        )
                        messages.success(request, 'Braille File Successfully Shared.')
                        any_new_entry_created = True

                if not any_new_entry_created:
                    messages.info(request, 'All Students in school year ' + str(school_year) + ' has been shared already with this file')
                return redirect('view_braille')
                
        braille_infos = BrailleInfo.objects.filter(user_id=user_id,deleteflag = False)

        usernames = UserProfile.objects.filter(
                is_student=True, 
                deleteflag=False)


        context = {'braille_infos': braille_infos,
                   'usernames': usernames,
                   'is_faculty': user_profile.is_faculty,
                   'school_years': unique_school_years
                   }
    else: 
        return redirect('login')
    return render(request, 'view_braille.html',context)


@login_required(login_url='login')
def download_braille(request, file_name):
    file_path = os.path.join('static/documents', file_name)
    print(file_path)

    # Check if the file exists before attempting to open it
    if not os.path.isfile(file_path):
        return HttpResponse(status=404)  # Return a 404 if the file is not found

    with open(file_path, 'rb') as doc_file:
        response = HttpResponse(
            doc_file.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{file_name}"'
        response['Cache-Control'] = 'no-cache'
        return response
    


@cache_control(no_cache=True, must_revalidate=True, no_store=True)
@login_required(login_url='login')
def archives(request):
    if request.user.is_authenticated:
        if request.method == 'POST':
            form_type = request.POST.get('form_type')
            if form_type == 'delete_braille':
                braille_id = request.POST.get('braille_id')
                try:
                    record_to_delete = BrailleInfo.objects.get(id=braille_id)
                    record_to_delete.delete()
                    messages.success(request, 'Deleted Successfully!')
                    return redirect('archives')
                except BrailleInfo.DoesNotExist:
                    messages.error(request, 'Braille is already Deleted!')
            elif form_type == 'restore_braille':
                braille_id = request.POST.get('braille_id')
                braille_info = BrailleInfo.objects.get(id=braille_id)
                braille_info.deleteflag = False
                braille_info.save()
                messages.success(request, 'Braille Successfully moved to Archive')
                return redirect('archives')
        user_id = request.user.id
        braille_infos = BrailleInfo.objects.filter(user_id=user_id,deleteflag = True)
        context = {'braille_infos': braille_infos}
        
    else:
        return redirect('login')
    
    return render(request, 'archives.html',context)




@csrf_exempt
def filter_students_disabled(request):
    if request.method == 'POST':
        school_year = request.POST.get('school_year')
        
        if school_year == "all" or not school_year:
            students = UserProfile.objects.filter(
                is_student=True, 
                is_active=False
            ).select_related('user').values(
                'user__id',
                'user__first_name',
                'user__last_name',
                'school_year',
                'user__username'
            )
        else:
            students = UserProfile.objects.filter(
                is_student=True, 
                is_active=False, 
                school_year=school_year
            ).select_related('user').values(
                'user__id',
                'user__first_name',
                'user__last_name',
                'school_year',
                'user__username'
            )
        return JsonResponse(list(students), safe=False)

@csrf_exempt
def filter_students_active(request):
    if request.method == 'POST':
        school_year = request.POST.get('school_year')
        
        if school_year == "all" or not school_year:
            students = UserProfile.objects.filter(
                is_student=True, 
                is_active = True
            ).select_related('user').values(
                'user__id',
                'user__first_name',
                'user__last_name',
                'school_year',
                'user__username'
            )
        else:
            students = UserProfile.objects.filter(
                is_student=True, 
                is_active = True,
                school_year=school_year
            ).select_related('user').values(
                'user__id',
                'user__first_name',
                'user__last_name',
                'school_year',
                'user__username'
            )
        return JsonResponse(list(students), safe=False)


@cache_control(no_cache=True, must_revalidate=True, no_store=True)
@login_required(login_url='login')
def list_of_student(request):
    if request.user.is_authenticated:
        if request.method == 'POST':
            form_type = request.POST.get('form_type')
            if form_type == 'disable_account':
                account_id = request.POST.get('account_id')
                print(account_id)
                user = User.objects.get(id=account_id)
                user_profile = UserProfile.objects.get(user = user)
                user_profile.is_active = False
                user_profile.save()
                messages.success(request, 'Account Disabled Successfully')

                return redirect('list_of_student')
            elif form_type == 'enable_account':
                
                account_id = request.POST.get('account_id')
                print(account_id)

                user = User.objects.get(id=account_id)
                user_profile = UserProfile.objects.get(user = user)
                user_profile.is_active = True
                user_profile.save()
                messages.success(request, 'Account Enabled Successfully')
                return redirect('list_of_student')
    students_active = UserProfile.objects.filter(is_student=True, is_active = True,deleteflag = False).select_related('user').values(
        'user__id',
        'user__first_name', 
        'user__last_name', 
        'school_year',
        'user__username'
    )
    students_disabled = UserProfile.objects.filter(is_student=True, is_active = False,deleteflag = False).select_related('user').values(
        'user__id',
        'user__first_name', 
        'user__last_name', 
        'school_year',
        'user__username'
    )
    
    return render(request, 'list_of_student.html', {'students_active': students_active, 'students_disaled': students_disabled})


@cache_control(no_cache=True, must_revalidate=True, no_store=True)
@login_required(login_url='login')
def manage_account(request):
    if request.user.is_authenticated:
        if request.method == 'POST':
            form_type = request.POST.get('form_type')
            if form_type == 'add_account':
                fname = request.POST.get('firstname')
                lname = request.POST.get('lastname')
                email = request.POST.get('email')
                role = request.POST.get('role')
                usernames = list(User.objects.values_list('username', flat=True))
                username = generate_unique_username(fname,lname,usernames)
                school_year = request.POST.get('school_year')
                pw = generate_random_string()
                email_check =  User.objects.filter(email = email).count()
                if email_check > 0:
                    messages.error(request,'Email Already Exists')
                else:
                    try:
                        validate_password(pw)
                        user = User.objects.create_user(username=username, email=email, password=pw, first_name=fname, last_name=lname)
                        user.save()
                        if role == 'teacher':
                            user_profile = UserProfile(
                                user=user,
                                initial_password=pw,
                                is_student=False,
                                is_faculty = True,
                                school_year = school_year
                            )
                            user_profile.save()
                        elif role == 'student':
                            user_profile = UserProfile(
                                user=user,
                                initial_password=pw,
                                is_student=True,
                                is_faculty = False,
                                school_year = school_year
                            )
                            user_profile.save()
                        fullname = fname + ' ' + lname
                        context = {
                            'fullname': fullname,
                            'username': username,
                            'initial_password': pw
                        }
                        
                        message_data = json.dumps(context)
                        messages.add_message(request, messages.SUCCESS, message_data, extra_tags='extra_info')
                        return redirect('manage_account')
                    except ValidationError as e:
                        # Handle the validation error (e.g., log it, display a message)
                        messages.error(request, f'Invalid: {", ".join(e.messages)}')
                profiles = UserProfile.objects.select_related('user').filter(deleteflag=False)
                context = {
                    'profiles':profiles,
                    'first_name': fname,
                    'last_name': lname,
                    'email': email
                }

                # return redirect('manage_account')
                return render(request, "manage_account.html", context)
            elif form_type == 'edit_account':
                user_id = request.POST.get('user_id')
                fname = request.POST.get('firstname')
                lname = request.POST.get('lastname')
                email = request.POST.get('email')
                print(user_id)
                user_profile = User.objects.get(id=user_id)
                user_profile.first_name = fname
                user_profile.last_name = lname
                user_profile.email = email
                user_profile.save()

                created_id = user_profile.id
                activity_history = ActivityHistory(user_id = user_id,activity_log="Edited Profile(User ID # " +str(created_id) + ")")
                activity_history.save()
                messages.success(request,'Account Updated Successfully!')
                return redirect('manage_account')

            elif form_type == 'delete_account':
                user_profile_id = request.POST.get('user_profile_id')
                user_profile = UserProfile.objects.get(id=user_profile_id)
                user_profile.deleteflag = True
                user_profile.save()
                messages.success(request,'Account Deleted Successfully!')
                return redirect('manage_account')

    # profiles = UserProfile.objects.select_related('user').filter(deleteflag=False, user_id != request.user.id)
    profiles = UserProfile.objects.select_related('user').filter(deleteflag=False).exclude(user_id=request.user.id)

    context = {
        'profiles':profiles
        }
    return render(request, 'manage_account.html',context)



#THIS IS A STUDENT ACCESS FOR SHARED BRAILLE

@cache_control(no_cache=True, must_revalidate=True, no_store=True)
@login_required(login_url='login')
def shared(request):
    if request.user.is_authenticated:
       user_id = request.user.id
       if request.method == 'POST':
            braille_id = request.POST.get('braille_id')
            filename = request.POST.get('filename')
            documents_dir = 'static/documents'
            
            fetch_braille = BrailleInfo.objects.get(id=braille_id, deleteflag=False)
            title = fetch_braille.title
            braille_text = fetch_braille.braille_text
            braille_draft = fetch_braille.braille_draft

            # Create the documents directory if it doesn't exist
            if not os.path.exists(documents_dir):
                os.makedirs(documents_dir)

            file_path = os.path.join(documents_dir, filename)

            # Check if the file exists, if not, recreate it
            if not os.path.exists(file_path):
                try:
                    # Create and save the document
                    # document = Document()
                    # document.add_heading(title, level=1)
                    # document.add_paragraph(braille_text)
                    # document.add_paragraph(braille_draft)
                    # document.save(file_path)

                    # Create a new document
                    document = Document()

                    # Add title with font size 15 and center alignment
                    title_paragraph = document.add_heading(title, level=1)
                    title_run = title_paragraph.runs[0]
                    title_run.font.size = Pt(15)
                    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centers the title

                    # Add space below the title (1 line break)
                    for _ in range(1):
                        document.add_paragraph()  # Add empty paragraphs for spacing

                    # Modify 'braille_text' to add 3 spaces between each character
                    braille_text_spaced = '   '.join(braille_text)  # Adds 3 spaces between each character

                    # Add braille_text with font size 22
                    para_braille = document.add_paragraph()
                    run_braille = para_braille.add_run(braille_text_spaced)
                    font_braille = run_braille.font
                    font_braille.size = Pt(22)

                    # Add braille_draft with font size 22
                    para_draft = document.add_paragraph()
                    run_draft = para_draft.add_run(braille_draft)
                    font_draft = run_draft.font
                    font_draft.size = Pt(22)

                    # Define the file path and save the document
                    file_path = os.path.join(documents_dir, filename)
                    document.save(file_path)
                except Exception as e:
                    print(f"Error creating document: {e}")

            try:
                # Determine the user's Downloads directory
                downloads_dir = os.path.join(os.path.expanduser("~"), "Downloads")
                destination_path = os.path.join(downloads_dir, filename)

                # Copy the file to the Downloads directory
                shutil.copy(file_path, destination_path)
                print(f"File copied to: {destination_path}")
                messages.success(request,"File Downloaded Successfully")
                # Optionally, you can also return a success message or redirect
                return redirect('shared')  # Adjust this to your success view
            except Exception as e:
                print(f"Error during file copy: {e}")
    else:
         return redirect('login')
   
    # Fetch SharedBraille entries where shared_to_user is the logged-in user
    shared_braille_entries = SharedBraille.objects.select_related('user', 'shared_to_user', 'braille_info').filter(shared_to_user=user_id)
    
    # Prepare the context dictionary correctly
    context = {'shared_braille_entries': shared_braille_entries}
    return render(request, 'shared.html',context)




#THIS IS AN ADMIN ACCESS FOR SHARED BRAILLE
@cache_control(no_cache=True, must_revalidate=True, no_store=True)
@login_required(login_url='login')
def shared_braille(request):
    if request.user.is_authenticated:
       user_id = request.user.id
       if request.method == 'POST':
            braille_id = request.POST.get('braille_id')
            filename = request.POST.get('filename')
            documents_dir = 'static/documents'
            
            fetch_braille = BrailleInfo.objects.get(id=braille_id, deleteflag=False)
            title = fetch_braille.title
            braille_text = fetch_braille.braille_text
            braille_draft = fetch_braille.braille_draft

            # Create the documents directory if it doesn't exist
            if not os.path.exists(documents_dir):
                os.makedirs(documents_dir)

            file_path = os.path.join(documents_dir, filename)

            # Check if the file exists, if not, recreate it
            if not os.path.exists(file_path):
                try:
                    # Create and save the document
                    # document = Document()
                    # document.add_heading(title, level=1)
                    # document.add_paragraph(braille_text)
                    # document.add_paragraph(braille_draft)
                    # document.save(file_path)

                    # Create a new document
                    document = Document()

                    # Add title with font size 15 and center alignment
                    title_paragraph = document.add_heading(title, level=1)
                    title_run = title_paragraph.runs[0]
                    title_run.font.size = Pt(15)
                    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Centers the title

                    # Add space below the title (1 line break)
                    for _ in range(1):
                        document.add_paragraph()  # Add empty paragraphs for spacing

                    # Modify 'braille_text' to add 3 spaces between each character
                    braille_text_spaced = '   '.join(braille_text)  # Adds 3 spaces between each character

                    # Add braille_text with font size 22
                    para_braille = document.add_paragraph()
                    run_braille = para_braille.add_run(braille_text_spaced)
                    font_braille = run_braille.font
                    font_braille.size = Pt(22)

                    # Add braille_draft with font size 22
                    para_draft = document.add_paragraph()
                    run_draft = para_draft.add_run(braille_draft)
                    font_draft = run_draft.font
                    font_draft.size = Pt(22)

                    # Define the file path and save the document
                    file_path = os.path.join(documents_dir, filename)
                    document.save(file_path)
                except Exception as e:
                    print(f"Error creating document: {e}")

            try:
                # Determine the user's Downloads directory
                downloads_dir = os.path.join(os.path.expanduser("~"), "Downloads")
                destination_path = os.path.join(downloads_dir, filename)

                # Copy the file to the Downloads directory
                shutil.copy(file_path, destination_path)
                print(f"File copied to: {destination_path}")
                messages.success(request,"File Downloaded Successfully")
                # Optionally, you can also return a success message or redirect
                return redirect('shared_braille')  # Adjust this to your success view
            except Exception as e:
                print(f"Error during file copy: {e}")
    else:
         return redirect('login')
   
    # Fetch SharedBraille entries where shared_to_user is the logged-in user
    shared_braille_entries = SharedBraille.objects.select_related('user', 'shared_to_user', 'braille_info').filter(user_id=user_id)
    
    # Prepare the context dictionary correctly
    context = {'shared_braille_entries': shared_braille_entries}
    return render(request, 'shared_braille.html',context)



@cache_control(no_cache=True, must_revalidate=True, no_store=True)
@login_required(login_url='login')
def account_settings(request):
    if request.user.is_authenticated:
       if request.method == 'POST':
            user_id = request.user.id
            old_password = request.POST.get('old_password')
            new_password1 = request.POST.get('new_password1')
            new_password2 = request.POST.get('new_password2')
            if not request.user.check_password(old_password):
                    messages.error(request, 'Old password is incorrect.')
                    return redirect('account_settings')
            # Check if new passwords match
            if new_password1 != new_password2:
                messages.error(request, 'New passwords do not match.')
                return redirect('account_settings')
            
            # Update user password
            request.user.set_password(new_password1)
            request.user.save()

            # save log every password change
            activity_history = ActivityHistory(user_id = user_id,activity_log='Password Changed')
            activity_history.save()
            # Update session to avoid auto-logout
            update_session_auth_hash(request, request.user)
            messages.success(request, 'Password changed successfully.')
            return redirect('logout_user')
       
    else:
         return redirect('login')
    return render(request, 'account_settings.html')





@cache_control(no_cache=True, must_revalidate=True, no_store=True)
def logout_user(request):
  logout(request)
  return redirect("login")